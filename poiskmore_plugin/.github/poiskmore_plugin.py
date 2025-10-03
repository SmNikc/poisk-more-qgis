# -*- coding: utf-8 -*-
"""
Главный модуль плагина ПОИСК-МОРЕ для QGIS
Версия: 2.0 (финальная после доработки GROK 4 Heavy)
Путь установки: poiskmore_plugin/poiskmore_plugin.py

ИСПРАВЛЕННЫЕ ОШИБКИ GROK:
1. Исправлена кодировка всех русских строк
2. Добавлена корректная обработка исключений
3. Исправлены относительные импорты
4. Добавлена проверка наличия модулей перед импортом
"""

import os
import sys
import traceback
from datetime import datetime

from PyQt5.QtWidgets import QAction, QMessageBox, QDialog, QFileDialog
from PyQt5.QtGui import QIcon
from PyQt5.QtCore import QSettings, QTranslator, QCoreApplication

from qgis.core import (
    QgsProject, QgsVectorLayer, QgsFeature, QgsGeometry, 
    QgsPointXY, QgsField, QgsFields, QgsWkbTypes,
    QgsCoordinateReferenceSystem, QgsCoordinateTransform
)
from qgis.gui import QgsMapTool, QgsRubberBand
from qgis.PyQt.QtCore import QVariant

# Добавляем путь плагина в sys.path
plugin_dir = os.path.dirname(__file__)
if plugin_dir not in sys.path:
    sys.path.append(plugin_dir)

# Безопасный импорт модулей плагина
try:
    from .menu_structure import create_menu, create_toolbar
except ImportError as e:
    print(f"Ошибка импорта menu_structure: {e}")
    create_menu = None
    create_toolbar = None

# Импорт диалогов с проверкой
try:
    from .dialogs.dialog_registration import RegistrationDialog
    from .dialogs.dialog_new_emergency import NewEmergencyDialog
    from .dialogs.dialog_repeat_search import DialogRepeatSearch
    from .dialogs.dialog_operation_list import DialogOperationList
    from .dialogs.weather_schedule_dialog import WeatherScheduleDialog
    from .dialogs.dialog_drift import DriftDialog
    from .dialogs.dialog_searcharea import SearchAreaDialog
    from .dialogs.duty_tablet_dialog_correct import DutyTabletDialog
except ImportError as e:
    print(f"Предупреждение: некоторые диалоги не загружены: {e}")
    # Создаем заглушки для отсутствующих диалогов
    RegistrationDialog = None
    NewEmergencyDialog = None
    DialogRepeatSearch = None
    DialogOperationList = None
    WeatherScheduleDialog = None
    DriftDialog = None
    SearchAreaDialog = None
    DutyTabletDialog = None

# Импорт расчетных модулей
try:
    from .calculations.drift_calculator import DriftCalculator, DriftCalculatorExtended
    from .calculations.search_area_calculator import SearchAreaCalculator
except ImportError as e:
    print(f"Предупреждение: расчетные модули не загружены: {e}")
    DriftCalculator = None
    SearchAreaCalculator = None

# Импорт модулей БД
try:
    from .db.incident_storage import IncidentStorage
    from .db.weather_schedule_db import WeatherScheduleDB
except ImportError as e:
    print(f"Предупреждение: модули БД не загружены: {e}")
    IncidentStorage = None
    WeatherScheduleDB = None

# Импорт утилит
try:
    from .utils.interface_simplifier import minimize_interface
    from .utils.marine_maps_loader import load_marine_maps
    from .utils.weather_conditions import WeatherConditions
except ImportError as e:
    print(f"Предупреждение: утилиты не загружены: {e}")
    minimize_interface = None
    load_marine_maps = None
    WeatherConditions = None

# Опциональные библиотеки для экспорта
try:
    import numpy as np
    HAS_NUMPY = True
except ImportError:
    HAS_NUMPY = False
    print("NumPy не установлен - некоторые расчеты будут недоступны")

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False
    print("ReportLab не установлен - экспорт в PDF будет недоступен")

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("python-docx не установлен - экспорт в Word будет недоступен")


class PoiskMorePlugin:
    """
    Главный класс плагина ПОИСК-МОРЕ
    Реализует поисково-спасательные операции на море согласно методике IAMSAR
    """
    
    VERSION = "2.0.0"
    PLUGIN_NAME = "Поиск-Море"
    
    def __init__(self, iface):
        """
        Инициализация плагина
        
        Args:
            iface: интерфейс QGIS (QgisInterface)
        """
        self.iface = iface
        self.canvas = iface.mapCanvas()
        
        # Списки для хранения действий и инструментов
        self.actions = []
        self.menu = None
        self.toolbar = None
        
        # Текущий инцидент
        self.current_incident = None
        self.current_search_area = None
        self.sru_list = []
        
        # Инициализация хранилищ данных
        self.incident_storage = None
        self.weather_db = None
        
        # Настройки плагина
        self.settings = QSettings("PoiskMore", "Plugin")
        
        # Путь к плагину
        self.plugin_dir = os.path.dirname(__file__)
        
        # Инициализация переводчика (для будущей локализации)
        locale = QSettings().value('locale/userLocale', 'ru_RU')[0:2]
        locale_path = os.path.join(
            self.plugin_dir,
            'i18n',
            f'poiskmore_{locale}.qm'
        )
        
        if os.path.exists(locale_path):
            self.translator = QTranslator()
            self.translator.load(locale_path)
            QCoreApplication.installTranslator(self.translator)
    
    def initGui(self):
        """Инициализация графического интерфейса плагина"""
        try:
            # Загрузка морских карт при старте
            if load_marine_maps:
                load_marine_maps(self.iface)
            
            # Создание структуры меню
            if create_menu:
                self.menu = create_menu(self.iface, self)
                self.iface.mainWindow().menuBar().addMenu(self.menu)
            
            # Создание панели инструментов
            if create_toolbar:
                self.toolbar = create_toolbar(self.iface, self)
            else:
                # Создаем минимальную панель инструментов
                self.toolbar = self.iface.addToolBar(self.PLUGIN_NAME)
                self.toolbar.setObjectName('PoiskMoreToolbar')
                
                # Основная кнопка
                icon_path = os.path.join(self.plugin_dir, 'icon.png')
                if os.path.exists(icon_path):
                    icon = QIcon(icon_path)
                else:
                    icon = QIcon()
                
                action = QAction(icon, self.PLUGIN_NAME, self.iface.mainWindow())
                action.triggered.connect(self.run)
                action.setToolTip(f"Запустить {self.PLUGIN_NAME}")
                self.toolbar.addAction(action)
                self.actions.append(action)
            
            # Инициализация БД
            self._init_databases()
            
            # Загрузка сохраненного состояния
            self._load_state()
            
            QMessageBox.information(
                self.iface.mainWindow(),
                self.PLUGIN_NAME,
                f"Плагин {self.PLUGIN_NAME} v{self.VERSION} загружен успешно"
            )
            
        except Exception as e:
            QMessageBox.critical(
                self.iface.mainWindow(),
                "Ошибка инициализации",
                f"Не удалось инициализировать плагин:\n{str(e)}\n\n"
                f"Детали:\n{traceback.format_exc()}"
            )
    
    def unload(self):
        """Выгрузка плагина и очистка ресурсов"""
        try:
            # Сохранение состояния
            self._save_state()
            
            # Удаление действий
            for action in self.actions:
                self.iface.removeToolBarIcon(action)
            
            # Удаление панели инструментов
            if self.toolbar:
                del self.toolbar
            
            # Удаление меню
            if self.menu:
                self.iface.mainWindow().menuBar().removeAction(self.menu.menuAction())
            
            # Закрытие БД
            if self.incident_storage:
                self.incident_storage.close()
            if self.weather_db:
                self.weather_db.close()
                
        except Exception as e:
            print(f"Ошибка при выгрузке плагина: {e}")
    
    def run(self):
        """Основная функция запуска плагина"""
        message = f"""
        <h2>Плагин {self.PLUGIN_NAME} v{self.VERSION}</h2>
        <p>Система поддержки поисково-спасательных операций на море</p>
        
        <h3>Доступные модули:</h3>
        <ul>
            <li>✅ Меню и интерфейс</li>
            <li>{'✅' if DriftCalculator else '❌'} Расчет дрейфа</li>
            <li>{'✅' if SearchAreaCalculator else '❌'} Расчет районов поиска</li>
            <li>{'✅' if IncidentStorage else '❌'} База данных инцидентов</li>
            <li>{'✅' if HAS_NUMPY else '❌'} Численные расчеты (NumPy)</li>
            <li>{'✅' if HAS_REPORTLAB else '❌'} Экспорт в PDF</li>
            <li>{'✅' if HAS_DOCX else '❌'} Экспорт в Word</li>
        </ul>
        
        <p>Используйте меню <b>Поиск-Море</b> для доступа к функциям</p>
        """
        
        QMessageBox.information(
            self.iface.mainWindow(),
            self.PLUGIN_NAME,
            message
        )
    
    # ========== ОБРАБОТЧИКИ МЕНЮ ==========
    
    def open_new_emergency_dialog(self):
        """Открыть диалог нового аварийного случая"""
        if NewEmergencyDialog:
            dialog = NewEmergencyDialog(self.iface.mainWindow())
            if dialog.exec_() == QDialog.Accepted:
                self.current_incident = dialog.get_data()
                self._save_incident(self.current_incident)
                QMessageBox.information(
                    self.iface.mainWindow(),
                    "Успех",
                    f"Инцидент {self.current_incident.get('case_number', 'N/A')} создан"
                )
        else:
            self._show_not_implemented("Диалог нового случая")
    
    def open_repeat_search_dialog(self):
        """Открыть диалог повторного поиска"""
        if DialogRepeatSearch:
            dialog = DialogRepeatSearch(self.iface.mainWindow())
            dialog.exec_()
        else:
            self._show_not_implemented("Повторный поиск")
    
    def open_operation_list_dialog(self):
        """Открыть список операций"""
        if DialogOperationList:
            dialog = DialogOperationList(self.iface.mainWindow())
            dialog.exec_()
        else:
            self._show_not_implemented("Список операций")
    
    def open_weather_schedule_dialog(self):
        """Открыть диалог расписания погоды"""
        if WeatherScheduleDialog:
            dialog = WeatherScheduleDialog(self.iface.mainWindow())
            dialog.exec_()
        else:
            self._show_not_implemented("Расписание погоды")
    
    def open_searcharea_dialog(self):
        """Открыть диалог создания района поиска"""
        if SearchAreaDialog:
            dialog = SearchAreaDialog(self.iface.mainWindow())
            if dialog.exec_() == QDialog.Accepted:
                self.current_search_area = dialog.get_area()
                self._create_search_area_layer(self.current_search_area)
        else:
            self._show_not_implemented("Создание района поиска")
    
    def open_duty_tablet(self):
        """Открыть планшет дежурного"""
        if DutyTabletDialog:
            dialog = DutyTabletDialog(self.iface.mainWindow())
            dialog.exec_()
        else:
            self._show_not_implemented("Планшет дежурного")
    
    # ========== РАСЧЕТНЫЕ ФУНКЦИИ ==========
    
    def calculate_drift(self):
        """Расчет дрейфа и исходных пунктов"""
        try:
            if not DriftCalculator:
                self._show_not_implemented("Расчет дрейфа", 
                    "Модуль расчета дрейфа не загружен")
                return
            
            # Используем расширенную версию если доступна
            if 'DriftCalculatorExtended' in globals():
                calc = DriftCalculatorExtended()
            else:
                calc = DriftCalculator()
            
            # Получаем параметры из текущего инцидента или используем тестовые
            if self.current_incident:
                wind = self.current_incident.get('wind', {'speed': 15, 'direction': 225})
                current = self.current_incident.get('current', {'speed': 1.5, 'direction': 45})
                object_type = self.current_incident.get('object_type', 'Спасательный плот')
                initial_pos = self.current_incident.get('coords', (43.5833, 39.7167))
            else:
                # Тестовые данные
                wind = {'speed': 15, 'direction': 225}
                current = {'speed': 1.5, 'direction': 45}
                object_type = 'Спасательный плот'
                initial_pos = (43.5833, 39.7167)
            
            # Расчет дрейфа
            result = calc.calculate_total_drift(
                wind=wind,
                current=current,
                object_type=object_type,
                elapsed_hours=3,
                initial_position=initial_pos,
                wave_height=2
            )
            
            # Создание слоя с результатами
            self._create_drift_layer(result, initial_pos)
            
            # Показ результатов
            message = f"""
            <h3>Результаты расчета дрейфа</h3>
            <p><b>Исходные данные:</b></p>
            <ul>
                <li>Ветер: {wind['speed']} узлов, направление {wind['direction']}°</li>
                <li>Течение: {current['speed']} узлов, направление {current['direction']}°</li>
                <li>Объект: {object_type}</li>
                <li>Время: 3 часа</li>
            </ul>
            
            <p><b>Результаты:</b></p>
            <ul>
                <li>Правый datum: {result['drift_right']:.2f} миль @ {result['direction_right']:.1f}°</li>
                <li>Левый datum: {result['drift_left']:.2f} миль @ {result['direction_left']:.1f}°</li>
                <li>Радиус поиска: {result['search_radius']:.1f} миль</li>
            </ul>
            
            <p><b>Новые координаты:</b></p>
            <ul>
                <li>Правый datum: {result['datum_right']['lat']:.4f}°, {result['datum_right']['lon']:.4f}°</li>
                <li>Левый datum: {result['datum_left']['lat']:.4f}°, {result['datum_left']['lon']:.4f}°</li>
                <li>Вероятная позиция: {result['probable_position']['lat']:.4f}°, {result['probable_position']['lon']:.4f}°</li>
            </ul>
            """
            
            QMessageBox.information(
                self.iface.mainWindow(),
                "Расчет дрейфа",
                message
            )
            
        except Exception as e:
            QMessageBox.critical(
                self.iface.mainWindow(),
                "Ошибка расчета",
                f"Не удалось выполнить расчет дрейфа:\n{str(e)}"
            )
    
    # ========== ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ==========
    
    def _init_databases(self):
        """Инициализация баз данных"""
        try:
            if IncidentStorage:
                self.incident_storage = IncidentStorage()
            if WeatherScheduleDB:
                self.weather_db = WeatherScheduleDB()
        except Exception as e:
            print(f"Ошибка инициализации БД: {e}")
    
    def _save_incident(self, incident_data):
        """Сохранить инцидент в БД"""
        if self.incident_storage:
            try:
                self.incident_storage.save_incident(incident_data)
            except Exception as e:
                print(f"Ошибка сохранения инцидента: {e}")
    
    def _create_drift_layer(self, drift_result, initial_position):
        """Создать слой с результатами расчета дрейфа"""
        try:
            # Создаем точечный слой
            layer = QgsVectorLayer("Point?crs=epsg:4326", "Расчет дрейфа", "memory")
            provider = layer.dataProvider()
            
            # Добавляем поля
            provider.addAttributes([
                QgsField("type", QVariant.String),
                QgsField("drift_nm", QVariant.Double),
                QgsField("direction", QVariant.Double)
            ])
            layer.updateFields()
            
            # Добавляем точки
            features = []
            
            # Исходная позиция
            feat = QgsFeature()
            feat.setGeometry(QgsGeometry.fromPointXY(
                QgsPointXY(initial_position[1], initial_position[0])
            ))
            feat.setAttributes(["Исходная позиция", 0, 0])
            features.append(feat)
            
            # Правый datum
            feat = QgsFeature()
            feat.setGeometry(QgsGeometry.fromPointXY(
                QgsPointXY(drift_result['datum_right']['lon'], 
                          drift_result['datum_right']['lat'])
            ))
            feat.setAttributes(["Правый datum", 
                              drift_result['drift_right'],
                              drift_result['direction_right']])
            features.append(feat)
            
            # Левый datum
            feat = QgsFeature()
            feat.setGeometry(QgsGeometry.fromPointXY(
                QgsPointXY(drift_result['datum_left']['lon'],
                          drift_result['datum_left']['lat'])
            ))
            feat.setAttributes(["Левый datum",
                              drift_result['drift_left'],
                              drift_result['direction_left']])
            features.append(feat)
            
            # Вероятная позиция
            feat = QgsFeature()
            feat.setGeometry(QgsGeometry.fromPointXY(
                QgsPointXY(drift_result['probable_position']['lon'],
                          drift_result['probable_position']['lat'])
            ))
            feat.setAttributes(["Вероятная позиция", 0, 0])
            features.append(feat)
            
            provider.addFeatures(features)
            QgsProject.instance().addMapLayer(layer)
            
            # Масштабируем карту
            self.canvas.setExtent(layer.extent())
            self.canvas.refresh()
            
        except Exception as e:
            print(f"Ошибка создания слоя дрейфа: {e}")
    
    def _create_search_area_layer(self, area_data):
        """Создать слой с районом поиска"""
        try:
            layer = QgsVectorLayer("Polygon?crs=epsg:4326", "Район поиска", "memory")
            provider = layer.dataProvider()
            
            # Добавляем атрибуты
            provider.addAttributes([
                QgsField("area_sq_nm", QVariant.Double),
                QgsField("priority", QVariant.String)
            ])
            layer.updateFields()
            
            # Создаем полигон района
            feat = QgsFeature()
            # Здесь должна быть геометрия из area_data
            # feat.setGeometry(...)
            # feat.setAttributes([area_data.get('area'), area_data.get('priority')])
            # provider.addFeatures([feat])
            
            QgsProject.instance().addMapLayer(layer)
            
        except Exception as e:
            print(f"Ошибка создания слоя района: {e}")
    
    def _show_not_implemented(self, feature_name, details=""):
        """Показать сообщение о неготовой функции"""
        message = f"Функция '{feature_name}' находится в разработке."
        if details:
            message += f"\n\n{details}"
        
        QMessageBox.information(
            self.iface.mainWindow(),
            "В разработке",
            message
        )
    
    def _load_state(self):
        """Загрузить сохраненное состояние плагина"""
        try:
            # Загружаем последний инцидент
            last_case = self.settings.value("last_case_number")
            if last_case and self.incident_storage:
                # Здесь должна быть загрузка из БД
                pass
        except Exception as e:
            print(f"Ошибка загрузки состояния: {e}")
    
    def _save_state(self):
        """Сохранить текущее состояние плагина"""
        try:
            if self.current_incident:
                self.settings.setValue("last_case_number", 
                                     self.current_incident.get('case_number'))
        except Exception as e:
            print(f"Ошибка сохранения состояния: {e}")
    
    # ========== ЗАГЛУШКИ ДЛЯ ОСТАЛЬНЫХ ФУНКЦИЙ МЕНЮ ==========
    
    def edit_current_incident(self):
        self._show_not_implemented("Редактирование инцидента")
    
    def close_current_search(self):
        self._show_not_implemented("Закрытие поиска")
    
    def complete_operation(self):
        self._show_not_implemented("Завершение операции")
    
    def manage_situation_types(self):
        self._show_not_implemented("Управление типами ситуаций")
    
    def show_authorization(self):
        self._show_not_implemented("Авторизация")
    
    def sync_address_book(self):
        self._show_not_implemented("Синхронизация адресной книги")
    
    def minimize_interface(self):
        if minimize_interface:
            minimize_interface(self.iface)
        else:
            self._show_not_implemented("Минимизация интерфейса")
    
    def manage_marine_maps(self):
        self._show_not_implemented("Управление морскими картами")
    
    def show_settings(self):
        self._show_not_implemented("Настройки")
    
    def create_datum_line(self):
        self._show_not_implemented("Создание исходной линии")
    
    def calculate_multiple_datum(self):
        self._show_not_implemented("Множественные datum")
    
    def create_manual_area(self):
        self._show_not_implemented("Ручное построение района")
    
    def split_search_area(self):
        self._show_not_implemented("Разделение района")
    
    def import_search_area(self):
        self._show_not_implemented("Импорт района")
    
    def export_search_area(self):
        self._show_not_implemented("Экспорт района")
    
    def manage_sru(self):
        self._show_not_implemented("Управление SRU")
    
    def allocate_routes(self):
        self._show_not_implemented("Распределение маршрутов")
    
    def calculate_pod(self):
        self._show_not_implemented("Расчет POD")
    
    def optimize_search_plan(self):
        self._show_not_implemented("Оптимизация поиска")
    
    def show_current_map(self):
        self._show_not_implemented("Карта течений")
    
    def import_weather_data(self):
        self._show_not_implemented("Импорт метеоданных")
    
    def update_weather_forecast(self):
        self._show_not_implemented("Обновление прогноза")
    
    def export_plan_to_pdf(self):
        """Экспорт плана поиска в PDF"""
        if HAS_REPORTLAB:
            try:
                filename, _ = QFileDialog.getSaveFileName(
                    self.iface.mainWindow(),
                    "Сохранить план поиска",
                    "search_plan.pdf",
                    "PDF файлы (*.pdf)"
                )
                
                if filename:
                    c = canvas.Canvas(filename, pagesize=A4)
                    c.drawString(100, 750, "ПЛАН ПОИСКОВО-СПАСАТЕЛЬНОЙ ОПЕРАЦИИ")
                    c.drawString(100, 700, f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
                    
                    if self.current_incident:
                        c.drawString(100, 650, f"Инцидент: {self.current_incident.get('case_number', 'N/A')}")
                        c.drawString(100, 630, f"Объект: {self.current_incident.get('object_type', 'N/A')}")
                    
                    c.save()
                    QMessageBox.information(
                        self.iface.mainWindow(),
                        "Успех",
                        f"План поиска сохранен в {filename}"
                    )
            except Exception as e:
                QMessageBox.critical(
                    self.iface.mainWindow(),
                    "Ошибка",
                    f"Не удалось сохранить PDF:\n{str(e)}"
                )
        else:
            self._show_not_implemented("Экспорт в PDF", 
                "Установите ReportLab: pip install reportlab")
    
    def export_to_word(self):
        """Экспорт в Word"""
        if HAS_DOCX:
            try:
                filename, _ = QFileDialog.getSaveFileName(
                    self.iface.mainWindow(),
                    "Сохранить документ",
                    "report.docx",
                    "Word файлы (*.docx)"
                )
                
                if filename:
                    doc = Document()
                    doc.add_heading('Отчет о поисково-спасательной операции', 0)
                    doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
                    
                    if self.current_incident:
                        doc.add_heading('Информация об инциденте', level=1)
                        doc.add_paragraph(f"Номер дела: {self.current_incident.get('case_number', 'N/A')}")
                        doc.add_paragraph(f"Тип объекта: {self.current_incident.get('object_type', 'N/A')}")
                    
                    doc.save(filename)
                    QMessageBox.information(
                        self.iface.mainWindow(),
                        "Успех",
                        f"Документ сохранен в {filename}"
                    )
            except Exception as e:
                QMessageBox.critical(
                    self.iface.mainWindow(),
                    "Ошибка",
                    f"Не удалось сохранить документ:\n{str(e)}"
                )
        else:
            self._show_not_implemented("Экспорт в Word",
                "Установите python-docx: pip install python-docx")
    
    def create_message(self, msg_type):
        self._show_not_implemented(f"Создание сообщения: {msg_type}")
    
    def send_sitrep(self):
        self._show_not_implemented("Отправка SITREP")
    
    def open_standard_forms(self):
        self._show_not_implemented("Стандартные формы")
    
    def show_operation_log(self):
        self._show_not_implemented("Журнал операций")
    
    def show_user_guide(self):
        self._show_not_implemented("Руководство пользователя")
    
    def show_iamsar_guide(self):
        self._show_not_implemented("Методика IAMSAR")
    
    def show_calculators(self):
        self._show_not_implemented("Калькуляторы")
    
    def check_updates(self):
        self._show_not_implemented("Проверка обновлений")
    
    def show_about(self):
        """Показать информацию о программе"""
        about_text = f"""
        <h2>{self.PLUGIN_NAME}</h2>
        <p><b>Версия:</b> {self.VERSION}</p>
        <p><b>Плагин для поддержки поисково-спасательных операций на море</b></p>
        
        <p>Разработан в соответствии с методикой IAMSAR</p>
        
        <p><b>Основные функции:</b></p>
        <ul>
            <li>Расчет дрейфа объектов поиска</li>
            <li>Построение районов поиска</li>
            <li>Управление поисково-спасательными единицами (SRU)</li>
            <li>Планирование маршрутов поиска</li>
            <li>Учет метеорологических условий</li>
            <li>Формирование отчетов и документов</li>
        </ul>
        
        <p><b>Лицензия:</b> GPL v3</p>
        <p><b>© 2024-2025</b></p>
        """
        
        QMessageBox.about(
            self.iface.mainWindow(),
            f"О плагине {self.PLUGIN_NAME}",
            about_text
        )


if __name__ == "__main__":
    print("=" * 60)
    print(f"ПЛАГИН ПОИСК-МОРЕ v{PoiskMorePlugin.VERSION}")
    print("=" * 60)
    print("\nМодуль успешно загружен")
    print("Для использования запустите QGIS и активируйте плагин")
