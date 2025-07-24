from docx import Document
import os
from datetime import datetime
from ..esb.esb_integration import send_message_via_esb

def generate_sitrep_docx(data):
    doc = Document()
    doc.add_heading("SITREP", 0)
    for label in ["type", "datetime", "sru", "coords"]:
        doc.add_paragraph(f"{label}: {data.get(label, 'N/A')}")
    doc.add_heading("Weather", level=1)
    doc.add_paragraph(data.get("weather", 'N/A'))
    doc.add_heading("Situation", level=1)
    doc.add_paragraph(data.get("situation", 'N/A'))
    doc.add_heading("Actions", level=1)
    doc.add_paragraph(data.get("actions", 'N/A'))
    doc.add_paragraph(f"Attachment: {data.get('attachment', 'N/A')}")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"SITREP_{timestamp}.docx"
    filepath = os.path.join(os.path.expanduser("~"), "Documents", filename)
    doc.save(filepath)
