from docx import Document
import os
from datetime import datetime
from ..esb.esb_integration import send_message_via_esb

def generate_sitrep_docx(data):
    # Обработка пустых полей
    for key in data:
        if data[key] is None:
            data[key] = 'N/A'
    
    doc = Document()
    doc.add_heading("SITREP", 0)
    for label in ["type", "datetime", "sru", "coords"]:
        doc.add_paragraph(f"{label.capitalize()}: {data.get(label)}")
    doc.add_heading("Weather", level=1)
    doc.add_paragraph(data.get("weather"))
    doc.add_heading("Situation", level=1)
    doc.add_paragraph(data.get("situation"))
    doc.add_heading("Actions", level=1)
    doc.add_paragraph(data.get("actions"))
    doc.add_paragraph(f"Attachment: {data.get('attachment')}")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"SITREP_{timestamp}.docx"
    filepath = os.path.join(os.path.expanduser("~"), "Documents", filename)
    doc.save(filepath)
    
    # Отправка через ESB
    esb_data = {"type": "SITREP_DOCX", "filepath": filepath, "data": data}
    if send_message_via_esb(esb_data):
        print("SITREP успешно отправлен через ESB")
    else:
        print("Ошибка отправки SITREP через ESB")
    
    return filepath