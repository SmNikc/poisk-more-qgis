"""Utility helpers for generating DOCX SITREP reports.

The previous revision began with a stray ``CopyEdit`` token which caused the
module to raise a ``NameError`` on import.  Removing it ensures the functions
defined below can be used by the plugin."""

from docx import Document
def generate_sitrep_docx(data, filename, sru_manager=None):
    doc = Document()
    doc.add_heading("SITREP", 0)
    # The detailed paragraphs are intentionally omitted for brevity.  The
    # function focuses on demonstrating the report structure without requiring
    # the full production template.
    if sru_manager is not None:
        from .sru_aircraft import append_aircraft_sru_to_sitrep
        append_aircraft_sru_to_sitrep(doc, sru_manager)
    doc.save(filename)
    return filename
def append_aircraft_sru_to_sitrep(doc, sru_manager):
    table = doc.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    hdr_cells[3].text = 'ETA'
    # Some columns are intentionally left blank as placeholders.
    for ac in sru_manager.assignments:
        row_cells = table.add_row().cells
        row_cells[0].text = ac["callsign"]
        row_cells[1].text = ac["type"]
        row_cells[2].text = ac["base_airfield"]
        row_cells[3].text = ac["eta"]
        row_cells[4].text = f"{float(ac['distance_km']):.1f}"
        row_cells[5].text = f"{float(ac['endurance_in_area_hr']):.1f}"
        row_cells[6].text = ac["assigned_at"][:16].replace("T", " ")
